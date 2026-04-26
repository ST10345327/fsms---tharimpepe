from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "Documentation"
GITHUB_URL = "https://github.com/ST10345327/fsms---tharimpepe"
DATE_TEXT = "20 April 2026"


@dataclass(frozen=True)
class TableSpec:
    title: str
    headers: list[str]
    rows: list[list[str]]


def set_page_margins(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    for style_name, size in [("Title", 20), ("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True

    if "Table Body" not in document.styles:
        style = document.styles.add_style("Table Body", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Aptos"
        style.font.size = Pt(9)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(9)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(document: Document, spec: TableSpec) -> None:
    if spec.title:
        p = document.add_paragraph()
        r = p.add_run(spec.title)
        r.bold = True
    table = document.add_table(rows=1, cols=len(spec.headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for i, header in enumerate(spec.headers):
        set_cell_text(header_cells[i], header, bold=True)
        shade_cell(header_cells[i], "D9EAF7")
    for row in spec.rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    document.add_paragraph()


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: Iterable[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_toc(document: Document) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "Right-click and update this table in Word to refresh page numbers."
    fld_sep.append(fld_text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])


def add_cover_page(document: Document, doc_title: str, subtitle: str) -> None:
    banner = document.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner_run = banner.add_run("THARIMPEPE FEEDING SCHEME")
    banner_run.bold = True
    banner_run.font.size = Pt(20)
    banner_run.font.color.rgb = RGBColor(255, 255, 255)
    shade_cell_like_paragraph(banner, "163A5F")

    document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("FEEDING SCHEME MANAGEMENT SYSTEM")
    title_run.bold = True
    title_run.font.size = Pt(18)

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(doc_title)
    sub_run.bold = True
    sub_run.font.size = Pt(15)

    document.add_paragraph()
    info = [
        f"Module: XISD5319 Work Integrated Learning",
        f"Task: {subtitle}",
        "Student: Olebogeng Phawe",
        "Student Number: ST10345327",
        "Client: Tharimpepe Feeding Scheme",
        "Location: Mafikeng, North West, South Africa",
        f"Date: {DATE_TEXT}",
        f"Implementation Repository: {GITHUB_URL}",
    ]
    for line in info:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run(
        "Logo area reserved for organisation branding and final lecturer submission formatting."
    ).italic = True

    document.add_page_break()


def shade_cell_like_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_footer(section, text: str) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = text


def paragraph(document: Document, text: str) -> None:
    document.add_paragraph(text)


def build_requirements_doc() -> Document:
    doc = Document()
    set_page_margins(doc)
    configure_styles(doc)
    add_cover_page(doc, "Requirement Analysis Document", "Task 2a - Requirement Analysis")
    add_footer(doc.sections[0], "XISD5319 | Task 2a | Requirement Analysis")

    doc.add_heading("Table of Contents", level=1)
    add_toc(doc)
    doc.add_page_break()

    doc.add_heading("1. Introduction / Problem Domain", level=1)
    paragraph(
        doc,
        "Tharimpepe Feeding Scheme is a community-based non-profit organisation that provides daily meals to underprivileged children and vulnerable community members in Mafikeng, North West. The organisation currently depends heavily on manual registers, paper attendance sheets, informal WhatsApp coordination, and ad hoc stock recording. These practices slow down service delivery, make records difficult to verify, and reduce the organisation’s ability to report accurately to donors and community partners.",
    )
    paragraph(
        doc,
        "The problem domain is defined by five operational pressures. First, beneficiary registration records are difficult to update and search when stored on paper. Second, daily attendance figures are vulnerable to omission and duplication, which affects meal planning. Third, food stock and donation information is not consistently captured in a central location, making it harder to monitor shortages, expiry dates, and donor contributions. Fourth, volunteer scheduling depends on informal communication, which creates confusion about who is available for each meal session. Fifth, reporting to donors and managers takes unnecessary time because information must be consolidated manually across separate notebooks and messages.",
    )
    paragraph(
        doc,
        "A web-based system is therefore needed to provide a central, secure, and structured way to manage beneficiary data, attendance, donations, stock, volunteer schedules, and management reports. This analysis document uses the assignment’s academic model as the primary presentation layer while noting, where relevant, that the implemented PHP/MySQL codebase contains richer implementation details such as additional reporting screens and support modules. Those supporting modules are treated as implementation details rather than the core scope of Task 2.",
    )

    doc.add_heading("2. Solution Domain", level=1)
    paragraph(
        doc,
        "The proposed solution is a web-based Feeding Scheme Management System built around a three-tier architecture: a browser-based presentation layer, a PHP application layer hosted on Apache, and a MySQL data layer. The system digitises beneficiary management, attendance recording, stock and donation tracking, volunteer scheduling, and management reporting while remaining within the project scope. Online payments, a mobile application, and advanced financial accounting remain out of scope.",
    )
    doc.add_heading("2.1 Functional Requirements", level=2)
    add_numbered(
        doc,
        [
            "The system must allow an Organisation Manager (Admin) to register, update, search, and deactivate beneficiary records.",
            "The system must allow authorised users to log in and log out securely.",
            "The system must record beneficiary attendance per meal session and date.",
            "The system must maintain food stock records, including quantities, units, expiry dates, and reorder levels.",
            "The system must capture donor profiles and donation records for food, cash, and other support.",
            "The system must allow volunteer records to be created, updated, and assigned to meal sessions.",
            "The system must generate attendance, stock, donor, volunteer, and dashboard summary reports.",
            "The system must display management information on a dashboard for operational monitoring.",
        ],
    )
    doc.add_heading("2.2 Non-Functional Requirements", level=2)
    add_numbered(
        doc,
        [
            "The user interface must be simple enough for staff and volunteers with limited digital literacy.",
            "The system must protect user credentials through hashed passwords and authenticated sessions.",
            "Normal page responses should remain efficient under the expected NGO workload of up to 500 beneficiaries and 50 users.",
            "The design must remain maintainable by following modular PHP/MVC-style separation and a normalized MySQL schema.",
            "The system must be compatible with common desktop browsers used by the organisation.",
            "The reporting design must support reliable, auditable outputs for managers, donors, and community partners.",
        ],
    )
    doc.add_heading("2.3 Use Case Catalogue", level=2)
    add_numbered(
        doc,
        [
            "Register beneficiary",
            "Log in",
            "Log out",
            "Record daily attendance",
            "Update beneficiary details",
            "Add food stock",
            "Add donation",
            "View stock levels",
            "Register volunteer",
            "Schedule volunteer",
            "View volunteer schedule",
            "Add donor",
            "Update donor",
            "Delete donor record",
            "Generate attendance report",
            "Generate stock report",
            "Generate donor report",
            "View dashboard summary",
        ],
    )
    paragraph(
        doc,
        "Implementation mapping note: the live repository realises these use cases through PHP controllers and models such as `BeneficiaryController`, `AttendanceController`, `FoodStockController`, `DonationController`, `VolunteerController`, `VolunteerScheduleController`, `ReportsController`, and `DashboardController`. The analysis model keeps a cleaner academic naming scheme for lecturer-facing clarity.",
    )

    use_case_rows = [
        ["Organisation Manager (Admin)", "Register beneficiary", "Beneficiary Registration Subsystem"],
        ["Organisation Manager (Admin)", "Log in", "User Account Feature"],
        ["Organisation Manager (Admin)", "Log out", "User Account Feature"],
        ["Volunteer", "Record daily attendance", "Attendance Subsystem"],
        ["Organisation Manager (Admin)", "Update beneficiary details", "Beneficiary Registration Subsystem"],
        ["Organisation Manager (Admin)", "Add food stock", "Stock & Donation Subsystem"],
        ["Organisation Manager (Admin)", "Add donation", "Stock & Donation Subsystem"],
        ["Volunteer", "View stock levels", "Stock & Donation Subsystem"],
        ["Organisation Manager (Admin)", "Register volunteer", "Volunteer Scheduling Subsystem"],
        ["Organisation Manager (Admin)", "Schedule volunteer", "Volunteer Scheduling Subsystem"],
        ["Volunteer", "View volunteer schedule", "Volunteer Scheduling Subsystem"],
        ["Organisation Manager (Admin)", "Add donor", "Stock & Donation Subsystem"],
        ["Organisation Manager (Admin)", "Update donor", "Stock & Donation Subsystem"],
        ["Organisation Manager (Admin)", "Delete donor record", "Stock & Donation Subsystem"],
        ["Organisation Manager (Admin)", "Generate attendance report", "Reporting Subsystem"],
        ["Organisation Manager (Admin)", "Generate stock report", "Reporting Subsystem"],
        ["Organisation Manager (Admin)", "Generate donor report", "Reporting Subsystem"],
        ["Organisation Manager (Admin)", "View dashboard summary", "Reporting Subsystem"],
    ]
    doc.add_heading("3. Use Case Table", level=1)
    add_table(
        doc,
        TableSpec(
            title="Table 1: Use Case Table",
            headers=["Participant (Active Actor)", "Function of the System", "Participant (Passive Actor)"],
            rows=use_case_rows,
        ),
    )

    logical_rows = [
        ["Username and password", "Dashboard or access error", "Authenticate user", "tbl_user"],
        ["Beneficiary details", "Beneficiary record confirmation", "Register beneficiary", "tbl_beneficiary"],
        ["Beneficiary ID and updated fields", "Updated beneficiary profile", "Update beneficiary details", "tbl_beneficiary"],
        ["Attendance date, meal session, presence status", "Attendance confirmation", "Record daily attendance", "tbl_attendance"],
        ["Food item details and quantity", "Updated stock list", "Add food stock", "tbl_food_stock"],
        ["Donor details and donation data", "Donation confirmation", "Add donation", "tbl_donor / tbl_donation"],
        ["Volunteer details", "Volunteer confirmation", "Register volunteer", "tbl_volunteer"],
        ["Volunteer ID, meal session, assignment details", "Schedule confirmation", "Schedule volunteer", "tbl_volunteer_schedule"],
        ["Date range / report type", "Attendance report", "Generate attendance report", "tbl_attendance / tbl_beneficiary / tbl_meal_session"],
        ["Date range / report type", "Stock report", "Generate stock report", "tbl_food_stock / tbl_donation"],
        ["Date range / report type", "Donor report", "Generate donor report", "tbl_donor / tbl_donation"],
        ["No new input", "Dashboard widgets and summaries", "View dashboard summary", "tbl_beneficiary / tbl_attendance / tbl_food_stock / tbl_donation"],
    ]
    doc.add_heading("4. Logical System Model Table", level=1)
    add_table(
        doc,
        TableSpec(
            title="Table 2: Logical System Model",
            headers=["GUI Input", "GUI Output", "System Process", "ER Table"],
            rows=logical_rows,
        ),
    )

    doc.add_heading("5. Classes", level=1)
    add_numbered(
        doc,
        [
            "User",
            "Beneficiary",
            "Volunteer",
            "Donor",
            "Donation",
            "FoodStock",
            "AttendanceRecord",
            "MealSession",
            "VolunteerSchedule",
            "Report",
        ],
    )

    business_rows = [
        ["User", "Beneficiary", "One-to-Many", "One admin can register and manage many beneficiaries."],
        ["User", "Volunteer", "One-to-Many", "One admin can manage many volunteer records."],
        ["Beneficiary", "AttendanceRecord", "One-to-Many", "One beneficiary can have many attendance records across meal sessions."],
        ["MealSession", "AttendanceRecord", "One-to-Many", "Each meal session can include many attendance records."],
        ["Volunteer", "VolunteerSchedule", "One-to-Many", "One volunteer can have multiple schedule assignments."],
        ["MealSession", "VolunteerSchedule", "One-to-Many", "A meal session can require multiple volunteer assignments."],
        ["Donor", "Donation", "One-to-Many", "One donor can make many donations over time."],
        ["FoodStock", "Donation", "One-to-Many", "One food stock item can be replenished by many food donations."],
        ["Report", "User", "Many-to-One", "Reports are generated by a requesting authorised user."],
    ]
    doc.add_heading("6. Business Rules Table", level=1)
    add_table(
        doc,
        TableSpec(
            title="Table 3: Business Rules",
            headers=["Entity", "Related Entity", "Relationship", "Business Rule & Notes"],
            rows=business_rows,
        ),
    )

    class_rows = [
        ["User", "UserID (int-11), Username (varchar-50), Email (varchar-100), PasswordHash (varchar-255), Role (varchar-20), IsActive (boolean)", "Beneficiary, Volunteer, Report"],
        ["Beneficiary", "BeneficiaryID (varchar-10), FullName (varchar-100), DateOfBirth (date), Address (varchar-255), GuardianName (varchar-100), ContactNumber (varchar-20), Status (varchar-20)", "AttendanceRecord, MealSession, User"],
        ["Volunteer", "VolunteerID (varchar-10), FullName (varchar-100), ContactNumber (varchar-20), Email (varchar-100), Skills (varchar-255), Status (varchar-20)", "VolunteerSchedule, MealSession, User"],
        ["Donor", "DonorID (varchar-10), DonorName (varchar-100), DonorType (varchar-30), ContactPerson (varchar-100), Phone (varchar-20), Email (varchar-100)", "Donation"],
        ["Donation", "DonationID (varchar-10), DonationDate (date), DonationType (varchar-30), ItemDescription (varchar-255), AmountZAR (decimal-10,2), Quantity (decimal-10,2)", "Donor, FoodStock, Report"],
        ["FoodStock", "FoodStockID (varchar-10), ItemName (varchar-100), Category (varchar-50), Unit (varchar-20), QuantityAvailable (decimal-10,2), ReorderLevel (decimal-10,2), ExpiryDate (date)", "Donation, Report"],
        ["AttendanceRecord", "AttendanceID (varchar-10), AttendanceDate (date), AttendanceStatus (varchar-10), Notes (varchar-255)", "Beneficiary, MealSession, Report"],
        ["MealSession", "MealSessionID (varchar-10), SessionDate (date), SessionType (varchar-30), Location (varchar-100), Notes (varchar-255)", "AttendanceRecord, VolunteerSchedule"],
        ["VolunteerSchedule", "ScheduleID (varchar-10), AssignedRole (varchar-50), ShiftStart (time), ShiftEnd (time), ScheduleStatus (varchar-20)", "Volunteer, MealSession"],
        ["Report", "ReportID (varchar-10), ReportName (varchar-100), ReportType (varchar-50), GeneratedDate (datetime), OutputFormat (varchar-20)", "User, AttendanceRecord, Donation, FoodStock"],
    ]
    doc.add_heading("7. Class Diagram Table", level=1)
    add_table(
        doc,
        TableSpec(
            title="Table 4: Class Diagram Table",
            headers=["Name of Entity (UML Class)", "Properties (type-size)", "Related to"],
            rows=class_rows,
        ),
    )

    doc.add_heading("8. Use Case Diagram", level=1)
    paragraph(doc, "[INSERT DIAGRAM HERE] Figure 1: Use Case Diagram - Feeding Scheme Management System")
    add_bullets(
        doc,
        [
            "Draw a system boundary labelled 'Feeding Scheme Management System'.",
            "Place actors outside the boundary: Organisation Manager (Admin), Volunteer, Donor, Beneficiary.",
            "Inside the boundary include use cases for login, logout, register beneficiary, update beneficiary, record attendance, add food stock, add donation, register volunteer, schedule volunteer, view volunteer schedule, generate attendance report, generate stock report, generate donor report, and view dashboard summary.",
            "Connect Organisation Manager to all management and reporting use cases; connect Volunteer to attendance, stock viewing, and schedule viewing; connect Donor to donation-related interaction; connect Beneficiary only where the lecturer expects a beneficiary-facing role or registration context.",
        ],
    )

    doc.add_heading("9. UML Class Diagram", level=1)
    paragraph(doc, "[INSERT DIAGRAM HERE] Figure 2: UML Class Diagram - Feeding Scheme Management System")
    add_bullets(
        doc,
        [
            "Show classes: User, Beneficiary, Volunteer, Donor, Donation, FoodStock, AttendanceRecord, MealSession, VolunteerSchedule, Report.",
            "Display key attributes from Table 4 inside each class box.",
            "Show multiplicities: User 1..* Beneficiary, User 1..* Volunteer, Beneficiary 1..* AttendanceRecord, MealSession 1..* AttendanceRecord, Volunteer 1..* VolunteerSchedule, MealSession 1..* VolunteerSchedule, Donor 1..* Donation, FoodStock 1..* Donation.",
            "Use simple association lines with multiplicity labels and keep implementation-only classes out of the academic diagram.",
        ],
    )

    doc.add_heading("10. Appendix", level=1)
    paragraph(
        doc,
        "Appendix A: Implementation reference. The live prototype and source code are available at the following GitHub repository for lecturer review and evidence of development work: "
        + GITHUB_URL,
    )
    paragraph(
        doc,
        "Appendix B: Academic-to-implementation mapping. The academic model uses `MealSession`, `AttendanceRecord`, `VolunteerSchedule`, and separate `Donor`/`Donation` entities for clarity and normalization. The implemented PHP/MySQL repository contains equivalent operational structures, but some tables and models use richer or slightly different names.",
    )

    doc.add_heading("11. Reference List", level=1)
    add_bullets(
        doc,
        [
            "Arola, M. (2022) Hardware Requirements for Web Development. Nashville Software School. Available at: https://learn.nashvillesoftwareschool.com/blog/ (Accessed: 20 April 2026).",
            "Daly, R. (2019) The pros and cons of switching to free and open-source digital tools, CultureHive. Available at: https://www.culturehive.co.uk (Accessed: 20 April 2026).",
            "Oracle (2025) MySQL 8.0 Reference Manual. Available at: https://dev.mysql.com/doc/ (Accessed: 20 April 2026).",
            "Satzinger, J.W., Jackson, R.B. and Burd, S.D. (2014) Systems Analysis and Design in a Changing World. 7th edn. Boston: Cengage Learning.",
            "Staman, O.D. (2004) Clear skies. PM Network, 18(3), pp. 40-46.",
            "Weierick, W. (2013) Three-Tier Architecture. Available at: https://www.tonymarston.net/php-mysql/3-tier-architecture.html (Accessed: 20 April 2026).",
        ],
    )
    return doc


def build_system_design_doc() -> Document:
    doc = Document()
    set_page_margins(doc)
    configure_styles(doc)
    add_cover_page(doc, "System Design Document", "Task 2b - System Design")
    add_footer(doc.sections[0], "XISD5319 | Task 2b | System Design")

    doc.add_heading("Table of Contents", level=1)
    add_toc(doc)
    doc.add_page_break()

    doc.add_heading("1. Introduction", level=1)
    paragraph(
        doc,
        "This System Design Document translates the approved requirements for the Tharimpepe Feeding Scheme Management System into a lecturer-aligned technical design. The system is a web-based management platform for beneficiary registration, attendance capture, food stock and donation control, volunteer scheduling, and operational reporting.",
    )
    paragraph(
        doc,
        "The design uses an academic model as its primary presentation layer and maps that model to the implemented repository where useful. The current implementation repository is available at "
        + GITHUB_URL
        + ". Supporting modules in the repository that are not central to the feeding scheme assignment scope are treated as secondary implementation detail rather than as core design entities.",
    )

    doc.add_heading("2. Logical Architectural Design - High-Level", level=1)
    paragraph(
        doc,
        "The Feeding Scheme Management System is designed as a three-tier web application. The Presentation Layer is accessed through a desktop or laptop web browser and provides HTML, CSS, and JavaScript interfaces for organisation managers and volunteers. The Application Layer runs on an Apache/PHP environment and contains the business logic, validation rules, routing, and report generation logic. The Data Layer consists of a MySQL relational database that stores normalized tables for users, beneficiaries, meal sessions, attendance, volunteers, donors, donations, food stock, schedules, and reports.",
    )
    paragraph(
        doc,
        "At runtime, a client submits an HTTP request through the browser. Apache forwards the request to the relevant PHP controller or route. PHP validates the session and business rules, reads from or writes to MySQL, then returns the resulting HTML view or report data to the browser. Functional building blocks are therefore divided as follows: input capture and display are handled at the client, business logic and access control are handled on the server, and persistent records are stored in the database.",
    )
    paragraph(doc, "[INSERT DIAGRAM HERE] Figure 1: High-Level Architecture Diagram")
    add_bullets(
        doc,
        [
            "Show a browser/client box on the left, Internet/HTTP in the middle, Apache/PHP application server next, and MySQL database on the right.",
            "Inside the application server note key modules: Authentication, Beneficiary Management, Attendance, Stock & Donations, Volunteer Scheduling, Reporting.",
            "Indicate request and response arrows between the browser and server, and query/result arrows between the server and database.",
        ],
    )

    doc.add_heading("3. Logical Architectural Design - Low-Level", level=1)
    paragraph(
        doc,
        "The low-level design links actors to the functions they perform and the database tables those functions use. The main actors are Organisation Manager (Admin), Volunteer, Donor, and Beneficiary. The main functions are expressed as use cases such as Register Beneficiary, Record Attendance, Add Donation, View Stock Levels, Schedule Volunteer, and Generate Reports. Each function reads from or writes to one or more normalized database tables.",
    )
    paragraph(
        doc,
        "The academic low-level design should group actors on the left, use cases in the center, and database tables on the right. Arrows must indicate which actor initiates which function and which function interacts with which table. This lecturer-facing diagram should remain simpler than the implemented codebase and focus on the core feeding-scheme workflow.",
    )
    paragraph(doc, "[INSERT DIAGRAM HERE] Figure 2: Low-Level Design Diagram")
    add_bullets(
        doc,
        [
            "Actors: Organisation Manager (Admin), Volunteer, Donor, Beneficiary.",
            "Functions: Login, Register Beneficiary, Update Beneficiary, Record Attendance, Add Food Stock, Add Donation, View Stock Levels, Register Volunteer, Schedule Volunteer, Generate Reports, View Dashboard.",
            "Database circles: tbl_user, tbl_beneficiary, tbl_attendance, tbl_meal_session, tbl_food_stock, tbl_donor, tbl_donation, tbl_volunteer, tbl_volunteer_schedule, tbl_report.",
            "Connect each function to the tables it reads or writes, for example Record Attendance -> tbl_attendance and tbl_meal_session; Add Donation -> tbl_donor and tbl_donation; View Dashboard -> multiple summary tables.",
        ],
    )

    doc.add_heading("4. User Interaction Design - Input Interactions", level=1)
    paragraph(
        doc,
        "Option 1 is used for input interaction design. The following hierarchical menu structure defines how users enter data into the system. These forms and menus support the identification of the attributes and services used in the class and database design.",
    )
    add_bullets(
        doc,
        [
            "Main Menu",
            "1. Beneficiary Management",
            "1.1 Register New Beneficiary - Full Name, Date of Birth, Address, Guardian Name, Contact Number, Status",
            "1.2 Update Beneficiary - Search by ID, edit profile fields, save changes",
            "1.3 Search/View Beneficiary - Search filters, beneficiary profile view",
            "2. Attendance",
            "2.1 Record Daily Attendance - Select Date, Select Meal Session, mark Present/Absent per beneficiary",
            "2.2 View Attendance - Filter by date, beneficiary, or meal session",
            "3. Food Stock & Donations",
            "3.1 Add Stock Entry - Item name, category, quantity, unit, expiry date, reorder level",
            "3.2 Add Donation - Donor, donation type, quantity or amount, linked stock item",
            "3.3 View Stock Levels - Search, filter, low-stock visibility",
            "4. Volunteer Management",
            "4.1 Register Volunteer - Name, contact, email, skills, status",
            "4.2 Assign Schedule - Volunteer, meal session, role, shift times",
            "4.3 View Schedules - Calendar or table view",
            "5. Donor Management",
            "5.1 Add Donor - Donor profile details",
            "5.2 Update Donor - Edit donor information",
            "5.3 View Donors - Search and list donors",
            "6. Reports & Dashboard",
            "6.1 Attendance Report parameters",
            "6.2 Stock/Donation Report parameters",
            "6.3 Volunteer Report parameters",
            "6.4 Donor Impact Report parameters",
        ],
    )
    paragraph(
        doc,
        "Input interaction description: each menu path leads to a form that collects validated data before it is processed by the application layer. Required fields, date controls, dropdown selections, and save or cancel actions are used to keep the interface simple and reliable for NGO operational users.",
    )

    doc.add_heading("5. User Interaction Design - Request Interactions", level=1)
    paragraph(
        doc,
        "Request interactions describe the service requests and outputs the user asks the system to produce. These interactions include search requests, report generation, schedule viewing, dashboard summaries, and printable output formats.",
    )
    add_bullets(
        doc,
        [
            "Main Request Menu",
            "1. Search Beneficiary -> enter search term -> results list on screen",
            "2. Generate Attendance Report -> choose date or date range -> screen table and PDF output",
            "3. View Stock Levels -> dashboard widget and detailed stock table",
            "4. Generate Stock/Donation Report -> select month or custom range -> screen and PDF output",
            "5. View Volunteer Schedule -> weekly or monthly calendar view",
            "6. Generate Volunteer Schedule Report -> selected period -> screen and PDF output",
            "7. Generate Donor Impact Report -> selected date range -> printable PDF summary",
            "8. View Dashboard Summary -> KPIs for beneficiaries, meals served, donations, stock alerts, and volunteer activity",
        ],
    )
    paragraph(
        doc,
        "Request interaction description: these outputs provide operational value rather than data capture. Each request is initiated through a menu or parameter form, processed by PHP business logic, and returned as an on-screen display, printable layout, or downloadable report.",
    )

    doc.add_heading("6. Database Design - Database Tables", level=1)
    paragraph(
        doc,
        "The database design is normalized to Third Normal Form (3NF). Each table stores a single subject area, non-key attributes depend on the primary key, and repeating or mixed donor-donation structures are separated into distinct entities for clarity. This academic design is intentionally cleaner than parts of the current implementation so that the lecturer-facing model remains consistent and normalized.",
    )
    paragraph(
        doc,
        "Implementation mapping note: the current repository includes some direct donation attributes on the donations structure and other implementation-specific reporting helpers. In this design, `tbl_donor` and `tbl_donation` are separated, and `tbl_meal_session` is introduced as a first-class academic table to support attendance and scheduling relationships cleanly.",
    )

    db_tables = [
        TableSpec(
            "Database Table #1: tbl_user",
            ["Primary Key", "Secondary Keys", "Data Fields"],
            [
                ["UserID", "Username, Email", "Username | Email | PasswordHash | Role | IsActive | CreatedAt"],
                ["U001", "admin.tharimpepe / admin@tharimpepe.org.za", "admin.tharimpepe | admin@tharimpepe.org.za | <hashed> | Admin | 1 | 2026-04-01"],
                ["U002", "volunteer1 / neo.molefe@gmail.com", "volunteer1 | neo.molefe@gmail.com | <hashed> | Volunteer | 1 | 2026-04-03"],
                ["U003", "manager.ops / ops@tharimpepe.org.za", "manager.ops | ops@tharimpepe.org.za | <hashed> | Admin | 1 | 2026-04-05"],
            ],
        ),
        TableSpec(
            "Database Table #2: tbl_beneficiary",
            ["Primary Key", "Secondary Keys", "Data Fields"],
            [
                ["BeneficiaryID", "RegisteredByUserID", "FullName | DOB | Address | GuardianName | ContactNo | Status | DateAdded"],
                ["B001", "U001", "Kgosi Moagi | 2015-03-12 | 12 Mmabatho St, Mafikeng | Mary Moagi | 0731234567 | Active | 2026-04-02"],
                ["B002", "U001", "Lesedi Tau | 2017-07-04 | 5 Tlokwe Ave, Mafikeng | John Tau | 0829876543 | Active | 2026-04-03"],
                ["B003", "U003", "Amahle Dube | 2016-11-20 | 8 Bophelo Rd, Mafikeng | Thandi Dube | 0761112233 | Active | 2026-04-04"],
            ],
        ),
        TableSpec(
            "Database Table #3: tbl_meal_session",
            ["Primary Key", "Secondary Keys", "Data Fields"],
            [
                ["MealSessionID", "SessionDate", "SessionType | Location | PlannedMeals | Notes"],
                ["MS001", "2026-04-10", "Breakfast | Tharimpepe Centre | 120 | Weekday morning session"],
                ["MS002", "2026-04-10", "Lunch | Tharimpepe Centre | 150 | School support meal"],
                ["MS003", "2026-04-11", "Lunch | Community Hall | 140 | Saturday outreach"],
            ],
        ),
        TableSpec(
            "Database Table #4: tbl_attendance",
            ["Primary Key", "Secondary Keys", "Data Fields"],
            [
                ["AttendanceID", "BeneficiaryID, MealSessionID", "AttendanceDate | AttendanceStatus | RecordedByUserID | Notes"],
                ["A001", "B001, MS001", "2026-04-10 | Present | U002 | Arrived on time"],
                ["A002", "B002, MS001", "2026-04-10 | Absent | U002 | Guardian reported illness"],
                ["A003", "B003, MS002", "2026-04-10 | Present | U001 | Lunch session"],
            ],
        ),
        TableSpec(
            "Database Table #5: tbl_volunteer",
            ["Primary Key", "Secondary Keys", "Data Fields"],
            [
                ["VolunteerID", "UserID, Email", "FullName | ContactNo | Email | Skills | Status | DateRegistered"],
                ["V001", "U002, neo.molefe@gmail.com", "Neo Molefe | 0712223344 | neo.molefe@gmail.com | Food prep, attendance | Active | 2026-04-03"],
                ["V002", "NULL, kabelo.setho@gmail.com", "Kabelo Setho | 0745558899 | kabelo.setho@gmail.com | Distribution | Active | 2026-04-04"],
                ["V003", "NULL, limpho.mokoena@gmail.com", "Limpho Mokoena | 0789123456 | limpho.mokoena@gmail.com | Stock handling | Active | 2026-04-05"],
            ],
        ),
        TableSpec(
            "Database Table #6: tbl_volunteer_schedule",
            ["Primary Key", "Secondary Keys", "Data Fields"],
            [
                ["ScheduleID", "VolunteerID, MealSessionID", "AssignedRole | ShiftStart | ShiftEnd | ScheduleStatus | Notes"],
                ["VS001", "V001, MS001", "Attendance Desk | 07:00 | 10:00 | Scheduled | Morning attendance support"],
                ["VS002", "V002, MS002", "Meal Distribution | 11:00 | 14:00 | Scheduled | Lunch service"],
                ["VS003", "V003, MS003", "Stock Control | 10:00 | 13:00 | Scheduled | Outreach stock issue"],
            ],
        ),
        TableSpec(
            "Database Table #7: tbl_donor",
            ["Primary Key", "Secondary Keys", "Data Fields"],
            [
                ["DonorID", "Email, Phone", "DonorName | DonorType | ContactPerson | Phone | Email | Address"],
                ["D001", "donations@ubuntufoods.co.za, 0183811000", "Ubuntu Foods | Organisation | Lerato Ndlovu | 0183811000 | donations@ubuntufoods.co.za | Mafikeng Industrial Area"],
                ["D002", "tshepo.khumalo@gmail.com, 0823344556", "Tshepo Khumalo | Individual | Tshepo Khumalo | 0823344556 | tshepo.khumalo@gmail.com | Mafikeng"],
                ["D003", "care@northwestpartners.org.za, 0183847712", "North West Partners | Organisation | Rina Jacobs | 0183847712 | care@northwestpartners.org.za | Mahikeng CBD"],
            ],
        ),
        TableSpec(
            "Database Table #8: tbl_food_stock",
            ["Primary Key", "Secondary Keys", "Data Fields"],
            [
                ["FoodStockID", "ItemName, Category", "ItemName | Category | QuantityAvailable | Unit | ReorderLevel | ExpiryDate | LastUpdated"],
                ["FS001", "Maize Meal, Grain", "Maize Meal | Grain | 180.00 | kg | 50.00 | 2026-08-30 | 2026-04-10"],
                ["FS002", "Cooking Oil, Pantry", "Cooking Oil | Pantry | 45.00 | litres | 15.00 | 2026-12-31 | 2026-04-10"],
                ["FS003", "Beans, Protein", "Beans | Protein | 95.00 | kg | 30.00 | 2026-09-15 | 2026-04-11"],
            ],
        ),
        TableSpec(
            "Database Table #9: tbl_donation",
            ["Primary Key", "Secondary Keys", "Data Fields"],
            [
                ["DonationID", "DonorID, FoodStockID", "DonationDate | DonationType | ItemDescription | Quantity | AmountZAR | RecordedByUserID"],
                ["DN001", "D001, FS001", "2026-04-08 | Food | Maize Meal Bags | 100.00 | NULL | U001"],
                ["DN002", "D002, NULL", "2026-04-09 | Cash | Cash Support | NULL | 2500.00 | U003"],
                ["DN003", "D003, FS003", "2026-04-10 | Food | Dry Beans | 40.00 | NULL | U001"],
            ],
        ),
        TableSpec(
            "Database Table #10: tbl_report",
            ["Primary Key", "Secondary Keys", "Data Fields"],
            [
                ["ReportID", "GeneratedByUserID, ReportType", "ReportName | ReportingPeriod | OutputFormat | GeneratedDate | StoragePath"],
                ["R001", "U001, Attendance", "Daily Attendance Report | 2026-04-10 | PDF | 2026-04-10 15:30 | /reports/attendance-2026-04-10.pdf"],
                ["R002", "U003, Stock", "Monthly Stock Report | 2026-04 | PDF | 2026-04-11 09:15 | /reports/stock-2026-04.pdf"],
                ["R003", "U001, Donor", "Donor Impact Report | 2026-Q2 | PDF | 2026-04-12 12:00 | /reports/donor-impact-q2.pdf"],
            ],
        ),
    ]
    for table in db_tables:
        add_table(doc, table)

    doc.add_heading("7. Database Design - ERD Diagrams", level=1)
    paragraph(doc, "[INSERT DIAGRAM HERE] ERD Diagram #1: Beneficiary - Attendance - MealSession")
    add_bullets(
        doc,
        [
            "Entities: tbl_beneficiary, tbl_attendance, tbl_meal_session.",
            "Relationships: one beneficiary to many attendance records; one meal session to many attendance records.",
            "Show PKs underlined and FKs in tbl_attendance referencing both tbl_beneficiary and tbl_meal_session.",
        ],
    )
    paragraph(doc, "[INSERT DIAGRAM HERE] ERD Diagram #2: Volunteer - VolunteerSchedule - MealSession")
    add_bullets(
        doc,
        [
            "Entities: tbl_volunteer, tbl_volunteer_schedule, tbl_meal_session.",
            "Relationships: one volunteer to many schedule records; one meal session to many schedule records.",
            "Use crow's foot notation and label cardinalities clearly.",
        ],
    )
    paragraph(doc, "[INSERT DIAGRAM HERE] ERD Diagram #3: Donor - Donation - FoodStock")
    add_bullets(
        doc,
        [
            "Entities: tbl_donor, tbl_donation, tbl_food_stock.",
            "Relationships: one donor to many donations; one food stock item can be linked to many food donations; cash donations may not link to stock.",
        ],
    )
    paragraph(doc, "[INSERT DIAGRAM HERE] ERD Diagram #4: Full System Overview")
    add_bullets(
        doc,
        [
            "Include all core tables: tbl_user, tbl_beneficiary, tbl_meal_session, tbl_attendance, tbl_volunteer, tbl_volunteer_schedule, tbl_donor, tbl_donation, tbl_food_stock, tbl_report.",
            "Label all PK and FK relationships and keep the layout uncluttered by grouping attendance, volunteer, and donor-stock subsystems.",
        ],
    )

    doc.add_heading("8. System Reports Design", level=1)
    report_rows = [
        ["Daily Attendance Report", "tbl_attendance, tbl_beneficiary, tbl_meal_session", "Screen + PDF", "Manager requests per date"],
        ["Monthly Stock Report", "tbl_food_stock, tbl_donation", "Screen + PDF", "Manager requests per month"],
        ["Donor Contribution Report", "tbl_donor, tbl_donation", "Printable PDF", "Manager or donor requests"],
        ["Volunteer Schedule Report", "tbl_volunteer, tbl_volunteer_schedule, tbl_meal_session", "Screen + PDF", "Manager requests per week"],
        ["Dashboard Summary", "tbl_beneficiary, tbl_attendance, tbl_food_stock, tbl_donation, tbl_volunteer_schedule", "Screen widgets", "System loads after login"],
    ]
    add_table(
        doc,
        TableSpec(
            title="Table 11: System Reports Design",
            headers=["Report Name", "Data Sources", "Output Format", "Trigger"],
            rows=report_rows,
        ),
    )
    paragraph(
        doc,
        "Report design note: the implemented repository already contains report-related controllers and views for attendance, donations, stock, schedules, and dashboard analytics. The academic design consolidates these into a smaller, normalized report model that remains consistent with the assignment brief.",
    )

    doc.add_heading("9. Appendix", level=1)
    paragraph(
        doc,
        "Appendix A: Implementation evidence is available at "
        + GITHUB_URL
        + " for lecturer inspection of the PHP/LAMP prototype, schema, controllers, models, and reporting views.",
    )
    paragraph(
        doc,
        "Appendix B: Diagram production note. Where final image files are not yet embedded, this document provides exact diagram content specifications so the figures can be produced in Draw.io and inserted without changing the surrounding analysis or numbering.",
    )

    doc.add_heading("10. Reference List", level=1)
    add_bullets(
        doc,
        [
            "Bootstrap (2024) Bootstrap Documentation. Available at: https://getbootstrap.com/docs/5.0/ (Accessed: 20 April 2026).",
            "Daly, R. (2019) The pros and cons of switching to free and open-source digital tools, CultureHive. Available at: https://www.culturehive.co.uk (Accessed: 20 April 2026).",
            "Oracle (2025) MySQL 8.0 Reference Manual. Available at: https://dev.mysql.com/doc/ (Accessed: 20 April 2026).",
            "PHP Group (2025) PHP Manual. Available at: https://www.php.net/manual/ (Accessed: 20 April 2026).",
            "Satzinger, J.W., Jackson, R.B. and Burd, S.D. (2014) Systems Analysis and Design in a Changing World. 7th edn. Boston: Cengage Learning.",
            "Weierick, W. (2013) Three-Tier Architecture. Available at: https://www.tonymarston.net/php-mysql/3-tier-architecture.html (Accessed: 20 April 2026).",
        ],
    )
    return doc


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    requirements_path = DOCS_DIR / "Task2a_Requirements_Analysis.docx"
    system_design_path = DOCS_DIR / "Task2b_System_Design.docx"

    build_requirements_doc().save(requirements_path)
    build_system_design_doc().save(system_design_path)

    print(f"Wrote {requirements_path}")
    print(f"Wrote {system_design_path}")


if __name__ == "__main__":
    main()
